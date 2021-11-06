from docx import Document
from docx.shared import Inches

class Doc(object):
    def __init__(self, filename) -> None:
        super().__init__()
        try:
            self.document = Document(filename)
        except BaseException:
            self.document = Document()
        self.filename = filename

    def add_text(self, text, *, style=None):
        if text != None:
            self.document.add_paragraph(text, style=style)

    def add_img(self, imgpath):
        self.document.add_picture(imgpath)

    def save(self, fileName=None):
        self.add_text("\n")
        if fileName != None:
            self.filename = fileName
        self.document.save(self.filename)
